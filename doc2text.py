"""
Extract text from .pdf / .docx / .doc for NLP; write JSON next to the source.

Output path: <same_dir>/<stem>.json (same stem as the input file).

PDF text uses the same extraction as pdfextract. DOCX uses python-docx.
Legacy .DOC is converted via LibreOffice (soffice) to DOCX (or PDF fallback), then extracted.

JSON shape (sumai-compatible pages list):
  {
    "source": "report.docx",
    "format": "docx",
    "pages": [{"page": 1, "text": "..."}],
    "text": "full concatenated text"
  }

Run: python doc2text.py -h
"""

from __future__ import annotations

import argparse
import json
import os
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

SUPPORTED = frozenset({".pdf", ".docx", ".doc"})


class NoTextLayerError(RuntimeError):
    """PDF (or converted doc) has pages but no extractable text layer."""


def _find_soffice() -> str | None:
    """Locate LibreOffice soffice executable."""
    for name in ("soffice", "soffice.exe", "libreoffice"):
        found = shutil.which(name)
        if found:
            return found
    if sys.platform == "win32":
        candidates = [
            Path(os.environ.get("PROGRAMFILES", r"C:\Program Files"))
            / "LibreOffice"
            / "program"
            / "soffice.exe",
            Path(os.environ.get("PROGRAMFILES(X86)", r"C:\Program Files (x86)"))
            / "LibreOffice"
            / "program"
            / "soffice.exe",
            Path(os.environ.get("LIBREOFFICE_PATH", "")) / "soffice.exe",
        ]
        for p in candidates:
            if p.is_file():
                return str(p)
    return None


def extract_pdf_pages(path: Path) -> list[tuple[int, str]]:
    """Reuse pdfextract page text helpers."""
    from pdfextract import _extract_page_text
    from pypdf import PdfReader

    with open(path, "rb") as f:
        reader = PdfReader(f)
        if getattr(reader, "is_encrypted", False) and reader.is_encrypted:
            # Try empty password (owner-only); otherwise fail clearly
            try:
                status = reader.decrypt("")
            except Exception:
                status = 0
            if status == 0:
                raise RuntimeError(
                    f"PDF is encrypted: {path.name}. Decrypt first (pdfdecrypt.py)."
                )
        n_pages = len(reader.pages)
        if n_pages == 0:
            raise RuntimeError(f"PDF has no pages: {path.name}")
        pages = []
        for i in range(n_pages):
            try:
                text = _extract_page_text(reader, i)
            except Exception:
                # Missing /Contents or broken page content stream → no text
                text = ""
            pages.append((i + 1, text))
        return pages


def pdf_has_text_layer(pages: list[tuple[int, str]]) -> bool:
    """True if any page has non-whitespace extractable text."""
    return any((t or "").strip() for _, t in pages)


def extract_docx_pages(path: Path) -> list[tuple[int, str]]:
    try:
        from docx import Document
    except ImportError as e:
        raise RuntimeError(
            "python-docx is required for .docx/.doc. "
            "Install: uv sync  (or pip install python-docx)"
        ) from e

    doc = Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            line = "\t".join(c for c in cells if c)
            if line:
                parts.append(line)
    text = "\n".join(parts).strip()
    return [(1, text)] if text else [(1, "")]


def _libreoffice_convert(src: Path, out_dir: Path, target: str) -> Path:
    """
    Convert src to target format (docx or pdf) in out_dir via soffice.
    Returns path to converted file.
    """
    soffice = _find_soffice()
    if not soffice:
        raise RuntimeError(
            "LibreOffice (soffice) not found. Install LibreOffice and add it to PATH, "
            "or set LIBREOFFICE_PATH to the LibreOffice program folder. "
            "Required for legacy .doc files."
        )
    out_dir.mkdir(parents=True, exist_ok=True)
    cmd = [
        soffice,
        "--headless",
        "--nologo",
        "--nofirststartwizard",
        "--convert-to",
        target,
        "--outdir",
        str(out_dir),
        str(src.resolve()),
    ]
    proc = subprocess.run(cmd, capture_output=True, text=True, timeout=300)
    if proc.returncode != 0:
        err = (proc.stderr or proc.stdout or "").strip()
        raise RuntimeError(f"LibreOffice convert failed ({proc.returncode}): {err[:500]}")

    # Output name is stem.target (LibreOffice uses original stem)
    converted = out_dir / f"{src.stem}.{target}"
    if not converted.is_file():
        # Sometimes LO lowercases or adjusts the name — search
        matches = list(out_dir.glob(f"{src.stem}*.{target}"))
        if not matches:
            matches = list(out_dir.glob(f"*.{target}"))
        if not matches:
            raise RuntimeError(
                f"LibreOffice reported success but no .{target} found in {out_dir}"
            )
        converted = matches[0]
    return converted


def extract_doc_pages(path: Path) -> list[tuple[int, str]]:
    """Legacy .doc: convert to .docx (preferred) or .pdf, then extract."""
    with tempfile.TemporaryDirectory(prefix="doc2text_") as tmp:
        tmp_path = Path(tmp)
        try:
            docx_path = _libreoffice_convert(path, tmp_path, "docx")
            return extract_docx_pages(docx_path)
        except RuntimeError as e1:
            try:
                pdf_path = _libreoffice_convert(path, tmp_path, "pdf")
                return extract_pdf_pages(pdf_path)
            except RuntimeError as e2:
                raise RuntimeError(
                    f"Could not extract .doc via LibreOffice.\n"
                    f"  docx attempt: {e1}\n"
                    f"  pdf attempt: {e2}"
                ) from e2


def pages_to_payload(
    source: Path,
    fmt: str,
    pages: list[tuple[int, str]],
) -> dict:
    page_objs = [{"page": p, "text": t} for p, t in pages]
    full = "\n\n".join(t for _, t in pages if t).strip()
    return {
        "source": source.name,
        "format": fmt,
        "pages": page_objs,
        "text": full,
        "has_text": bool(full),
        "page_count": len(pages),
    }


def extract_to_json(
    input_path: Path,
    output_path: Path | None = None,
    *,
    allow_empty: bool = False,
) -> Path:
    """
    Extract text from input_path and write JSON beside it (or to output_path).
    Returns the JSON path written.

    For PDFs with no extractable text layer, raises NoTextLayerError unless
    allow_empty=True (still writes JSON with has_text=false).
    """
    path = input_path.resolve()
    if not path.is_file():
        raise FileNotFoundError(f"Not a file: {path}")
    ext = path.suffix.lower()
    if ext not in SUPPORTED:
        raise ValueError(f"Unsupported type {ext!r}; expected {sorted(SUPPORTED)}")

    if ext == ".pdf":
        pages = extract_pdf_pages(path)
        fmt = "pdf"
        if not pdf_has_text_layer(pages):
            msg = (
                f"No text layer in PDF: {path.name} "
                f"({len(pages)} page(s); likely scanned/image-only). "
                f"Run OCR first: python pdfocr.py -i \"{path}\" -O"
            )
            if not allow_empty:
                raise NoTextLayerError(msg)
            print(f"Warning: {msg}", file=sys.stderr)
    elif ext == ".docx":
        pages = extract_docx_pages(path)
        fmt = "docx"
    else:
        pages = extract_doc_pages(path)
        fmt = "doc"
        # .doc converted via PDF may also lack a text layer
        if fmt == "doc" and not any((t or "").strip() for _, t in pages):
            # Could be empty Word doc OR pdf-fallback with no text — warn/fail for PDF-like emptiness
            # Only treat as no-text-layer if we got multiple empty "pages" from PDF path is hard to know.
            # If single empty page from docx path, that's a blank document (allow write).
            pass

    if not any(t for _, t in pages) and fmt != "pdf":
        print(
            f"Note: No text extracted from {path.name} (empty document).",
            file=sys.stderr,
        )

    out = output_path if output_path is not None else path.with_suffix(".json")
    out = out.resolve()
    payload = pages_to_payload(path, fmt, pages)
    out.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return out


def iter_inputs(folder: Path, recursive: bool) -> list[Path]:
    pattern = "**/*" if recursive else "*"
    files = []
    for p in sorted(folder.glob(pattern)):
        if p.is_file() and p.suffix.lower() in SUPPORTED:
            files.append(p)
    return files


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Extract text from PDF/DOCX/DOC to same-stem JSON for NLP.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python doc2text.py -i report.docx          # -> report.json beside it
  python doc2text.py -i scan.pdf
  python doc2text.py -i legacy.doc           # needs LibreOffice
  python doc2text.py -d ./inbox -r
  python doc2text.py -i report.docx -o elsewhere/out.json

JSON includes pages[] (sumai-compatible) and a concatenated text field.
        """,
    )
    parser.add_argument("-i", "--input", help="Input .pdf / .docx / .doc file")
    parser.add_argument(
        "-d",
        "--directory",
        metavar="DIR",
        help="Process all supported files in DIR",
    )
    parser.add_argument(
        "-r",
        "--recursive",
        action="store_true",
        help="With -d: include subdirectories",
    )
    parser.add_argument(
        "-o",
        "--output",
        default=None,
        help="Output JSON path (single-file mode only; default: <stem>.json beside input)",
    )
    parser.add_argument(
        "--skip-existing",
        action="store_true",
        help="Skip if destination .json already exists",
    )
    parser.add_argument(
        "--allow-empty",
        action="store_true",
        help="Write JSON even when a PDF has no text layer (default: fail with OCR hint)",
    )
    parser.add_argument("-v", "--verbose", action="store_true")
    args = parser.parse_args()

    if bool(args.input) == bool(args.directory):
        parser.error("Provide exactly one of -i/--input or -d/--directory")

    if args.directory:
        if args.output:
            parser.error("-o/--output is only valid with -i/--input")
        dir_path = Path(args.directory)
        if not dir_path.is_dir():
            print(f"Error: Not a directory: {args.directory}", file=sys.stderr)
            sys.exit(1)
        files = iter_inputs(dir_path, args.recursive)
        if not files:
            print(f"No .pdf/.docx/.doc files found in {dir_path}", file=sys.stderr)
            sys.exit(1)
        ok = 0
        for f in files:
            dest = f.with_suffix(".json")
            if args.skip_existing and dest.is_file():
                if args.verbose:
                    print(f"Skip existing {dest}")
                ok += 1
                continue
            try:
                out = extract_to_json(f, allow_empty=args.allow_empty)
                print(f"Wrote {out}")
                ok += 1
            except Exception as e:
                print(f"Error {f}: {e}", file=sys.stderr)
        sys.exit(0 if ok == len(files) else 1)

    inp = Path(args.input)
    dest = Path(args.output) if args.output else inp.with_suffix(".json")
    if args.skip_existing and dest.is_file():
        print(f"Skip existing {dest}")
        sys.exit(0)
    try:
        out = extract_to_json(
            inp,
            dest if args.output else None,
            allow_empty=args.allow_empty,
        )
        print(f"Wrote {out}")
        sys.exit(0)
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
